# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright 2026 Shawn Hartsock and contributors

"""MarkdownToDocx — walk a mistune 3.x AST and emit a python-docx Document."""

from __future__ import annotations

import io
import os
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt

from .mermaid import MermaidUnavailable, render_mermaid_to_png

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


class MarkdownToDocx:
    """Convert a Markdown string to a python-docx Document.

    Args:
        source: Markdown source text.
        render_diagrams: When True (default), Mermaid fenced blocks are
            rendered to PNG via mmdc and embedded as images. When False
            (or when mmdc is unavailable), the fenced block is kept as a
            monospace code block.
    """

    def __init__(self, source: str, *, render_diagrams: bool = True) -> None:
        self.source = source
        self.render_diagrams = render_diagrams

    def build(self) -> Document:
        import mistune  # deferred so the module can be imported for type-checking

        doc = Document()
        md = mistune.create_markdown(renderer="ast", plugins=["table"])
        tokens = md(self.source) or []
        _BlockVisitor(doc, render_diagrams=self.render_diagrams).visit_all(tokens)
        return doc


# ---------------------------------------------------------------------------
# Internal AST visitors
# ---------------------------------------------------------------------------


class _BlockVisitor:
    def __init__(self, doc: Document, *, render_diagrams: bool) -> None:
        self.doc = doc
        self.render_diagrams = render_diagrams

    def visit_all(self, tokens: list[dict]) -> None:
        for token in tokens:
            self._visit(token)

    def _visit(self, token: dict) -> None:
        t = token.get("type", "")
        handler = getattr(self, f"_on_{t}", None)
        if handler:
            handler(token)

    # --- block handlers ---

    def _on_heading(self, token: dict) -> None:
        level = token.get("attrs", {}).get("level", 1)
        text = _inline_text(token.get("children", []))
        self.doc.add_heading(text, level=level)

    def _on_paragraph(self, token: dict) -> None:
        children = token.get("children", [])
        # A lone image in a paragraph gets its own picture paragraph.
        if len(children) == 1 and children[0].get("type") == "image":
            self._embed_image(children[0])
            return
        p = self.doc.add_paragraph()
        _InlineVisitor(p).visit_all(children)

    def _on_block_code(self, token: dict) -> None:
        info = (token.get("attrs") or {}).get("info") or ""
        lang = info.split()[0] if info else ""
        # mistune's block_code raw includes a trailing newline; strip it so
        # the rendered paragraph text matches the source faithfully.
        code = token.get("raw", "").rstrip("\n")

        if lang == "mermaid" and self.render_diagrams:
            self._render_mermaid(code)
            return

        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def _on_list(self, token: dict) -> None:
        ordered = (token.get("attrs") or {}).get("ordered", False)
        for item in token.get("children", []):
            self._on_list_item(item, ordered=ordered)

    def _on_list_item(self, token: dict, *, ordered: bool = False) -> None:
        style = "List Number" if ordered else "List Bullet"
        p = self.doc.add_paragraph(style=style)
        for child in token.get("children", []):
            ct = child.get("type", "")
            if ct in ("paragraph", "block_text"):
                # mistune emits `block_text` for list-item content in tight
                # lists and `paragraph` in loose lists; both wrap inline runs.
                _InlineVisitor(p).visit_all(child.get("children", []))
            else:
                # Nested block inside list item — recurse.
                self._visit(child)

    def _on_block_quote(self, token: dict) -> None:
        self.visit_all(token.get("children", []))

    def _on_table(self, token: dict) -> None:
        children = token.get("children", [])
        head_tok = next((c for c in children if c.get("type") == "table_head"), None)
        body_tok = next((c for c in children if c.get("type") == "table_body"), None)

        # mistune's table plugin puts table_cells directly inside table_head
        # (no intermediate row wrapper) but wraps body cells in table_row.
        # Normalize so head and body have the same shape: list[row] where
        # row.children == list[cell].
        head_rows = [head_tok] if head_tok else []
        body_rows = (body_tok or {}).get("children", [])
        all_rows = head_rows + body_rows
        if not all_rows:
            return

        num_cols = len(all_rows[0].get("children", []))
        if num_cols == 0:
            return

        table = self.doc.add_table(rows=len(all_rows), cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row_tok in enumerate(all_rows):
            for col_idx, cell_tok in enumerate(row_tok.get("children", [])):
                if col_idx < num_cols:
                    table.cell(row_idx, col_idx).text = _inline_text(
                        cell_tok.get("children", [])
                    )

    def _on_thematic_break(self, token: dict) -> None:  # noqa: ARG002
        self.doc.add_paragraph("─" * 40)

    def _on_blank_line(self, token: dict) -> None:  # noqa: ARG002
        pass

    # --- helpers ---

    def _embed_image(self, token: dict) -> None:
        url = (token.get("attrs") or {}).get("url", "")
        alt = (token.get("attrs") or {}).get("alt", url)
        if url and os.path.exists(url):
            try:
                self.doc.add_picture(url, width=Inches(5.5))
                return
            except Exception:
                pass
        self.doc.add_paragraph(f"[Image: {alt}]")

    def _render_mermaid(self, source: str) -> None:
        try:
            png = render_mermaid_to_png(source)
            png = _embed_mermaid_source(png, source)
            self.doc.add_picture(io.BytesIO(png), width=Inches(5.5))
        except MermaidUnavailable:
            # Fall back to monospace block.
            p = self.doc.add_paragraph()
            run = p.add_run(source)
            run.font.name = "Courier New"
            run.font.size = Pt(9)


class _InlineVisitor:
    """Adds formatted runs to an existing paragraph."""

    def __init__(self, paragraph: Paragraph) -> None:
        self.p = paragraph

    def visit_all(self, tokens: list[dict]) -> None:
        for token in tokens:
            self._visit(token, bold=False, italic=False)

    def _visit(self, token: dict, *, bold: bool, italic: bool) -> None:
        t = token.get("type", "")
        if t in ("text", "raw"):
            self._run(token.get("raw", ""), bold=bold, italic=italic)
        elif t == "strong":
            for child in token.get("children", []):
                self._visit(child, bold=True, italic=italic)
        elif t == "emphasis":
            for child in token.get("children", []):
                self._visit(child, bold=bold, italic=True)
        elif t == "codespan":
            run = self._run(token.get("raw", ""), bold=bold, italic=italic)
            run.font.name = "Courier New"
        elif t == "link":
            # Render link text; URL is lost (Word hyperlinks require COM/OPC manipulation).
            for child in token.get("children", []):
                self._visit(child, bold=bold, italic=italic)
        elif t in ("softline", "softbreak", "linebreak"):
            self._run(" ", bold=False, italic=False)
        elif t == "image":
            alt = (token.get("attrs") or {}).get("alt", "")
            self._run(f"[{alt}]", bold=False, italic=False)

    def _run(self, text: str, *, bold: bool, italic: bool):
        run = self.p.add_run(text)
        run.bold = bold
        run.italic = italic
        return run


def _embed_mermaid_source(png: bytes, source: str) -> bytes:
    """Embed the Mermaid source into the PNG's iTXt metadata, round-trippable
    via `scrybe_mermaid.extract`. Falls back to the unmodified PNG if the
    `scrybe_mermaid` binding (or its embed step) is unavailable, so export
    never fails just because the codec is missing.
    """
    try:
        import scrybe_mermaid

        return scrybe_mermaid.embed(png, source)
    except Exception:
        return png


def _inline_text(tokens: list[dict]) -> str:
    """Extract plain text from inline tokens (no formatting).

    mistune 3.x emits inline text as `{"type": "text", "raw": "..."}`.
    Earlier versions used `"type": "raw"` — accept both for resilience.
    """
    parts: list[str] = []
    for tok in tokens:
        t = tok.get("type", "")
        if t in ("text", "raw"):
            parts.append(tok.get("raw", ""))
        elif t in ("strong", "emphasis", "link"):
            parts.append(_inline_text(tok.get("children", [])))
        elif t == "codespan":
            parts.append(tok.get("raw", ""))
        elif t in ("softline", "softbreak", "linebreak"):
            parts.append(" ")
        elif t == "image":
            parts.append((tok.get("attrs") or {}).get("alt", ""))
    return "".join(parts)
